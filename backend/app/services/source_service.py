from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse

import docx
import requests
from bs4 import BeautifulSoup
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models import Agent, AgentRevision, KnowledgeChunk, KnowledgeDocument
from app.utils.text import chunk_text, normalize_whitespace, sha256_text, tokenize


settings = get_settings()


FALLBACK_KNOWLEDGE = [
    {
        "title": "Checking your Atome Card application status",
        "source_url": "https://help.atome.ph/hc/en-gb/articles/8712978836377-How-can-I-check-the-status-of-my-application",
        "content": (
            "Customers can check their Atome Card application status in the Atome app. "
            "The help center article explains where the application progress is shown and directs users to follow in-app updates."
        ),
    },
    {
        "title": "Why an Atome Card transaction can fail",
        "source_url": "https://help.atome.ph/hc/en-gb/articles/43466116960153-Why-did-my-card-transaction-fail",
        "content": (
            "Atome Card transactions may fail because of merchant restrictions, insufficient available limit, incorrect card details, risk checks, or network issues. "
            "Customers should verify the merchant, card status, and spending availability before retrying."
        ),
    },
    {
        "title": "Pending Atome Card transaction guidance",
        "source_url": "https://help.atome.ph/hc/en-gb/articles/42056424762905-What-should-I-do-if-my-transaction-status-shows-Pending",
        "content": (
            "If an Atome Card transaction status shows Pending, the payment process is not complete yet. "
            "If the charge is for an unsuccessful transaction, it will typically be automatically cancelled or refunded within 14 days."
        ),
    },
]


@dataclass
class ParsedDocument:
    title: str
    content: str
    source_type: str
    source_url: str | None = None
    filename: str | None = None
    mime_type: str | None = None
    payload_json: dict[str, Any] = field(default_factory=dict)


@dataclass
class SyncOutcome:
    documents_synced: int
    chunks_synced: int
    sync_mode: str
    fallback_used: bool
    last_sync_warning: str | None = None


class SourceService:
    def __init__(self) -> None:
        self.session = requests.Session()
        self.session.headers.update(
            {"User-Agent": "atome-chatbot-interview-demo/1.0 (+https://help.atome.ph)"}
        )

    def build_system_prompt(self, agent_name: str) -> str:
        return (
            f"You are {agent_name}, a careful customer service assistant. "
            "Answer from verified knowledge only, cite the retrieved sources, "
            "and route case-specific account questions into the correct tools."
        )

    def _extract_help_center_config(self, category_url: str) -> tuple[str, str]:
        parsed = urlparse(category_url)
        path_parts = [part for part in parsed.path.split("/") if part]

        locale = "en-gb"
        if "hc" in path_parts:
            hc_index = path_parts.index("hc")
            if hc_index + 1 < len(path_parts):
                locale = path_parts[hc_index + 1]

        category_id = ""
        for index, part in enumerate(path_parts):
            if part == "categories" and index + 1 < len(path_parts):
                next_part = path_parts[index + 1]
                category_id = next_part.split("-", 1)[0]
                break
            if part.isdigit():
                category_id = part
            elif "-" in part and part.split("-", 1)[0].isdigit():
                category_id = part.split("-", 1)[0]

        if not category_id:
            query_category = parse_qs(parsed.query).get("category")
            if query_category:
                category_id = query_category[0]

        if not category_id:
            raise ValueError(f"Could not parse category ID from {category_url}")
        return locale, category_id

    def _fetch_json(self, url: str) -> dict[str, Any]:
        response = self.session.get(url, timeout=settings.request_timeout_seconds)
        response.raise_for_status()
        return response.json()

    def _fetch_paginated_items(self, initial_url: str, item_key: str) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        next_url: str | None = initial_url
        while next_url:
            payload = self._fetch_json(next_url)
            items.extend(payload.get(item_key, []))
            next_url = payload.get("next_page")
        return items

    def _html_to_text(self, html: str) -> str:
        soup = BeautifulSoup(html or "", "html.parser")
        return normalize_whitespace(soup.get_text(" ", strip=True))

    def _fallback_documents(self) -> list[ParsedDocument]:
        return [
            ParsedDocument(
                source_type="kb_article",
                payload_json={
                    "sync_mode": "fallback",
                    "section_name": "Fallback knowledge",
                    "provenance": "seed_fallback",
                },
                **item,
            )
            for item in FALLBACK_KNOWLEDGE
        ]

    def parse_public_help_center(self, category_url: str) -> list[ParsedDocument]:
        parsed_base = urlparse(category_url)
        locale, category_id = self._extract_help_center_config(category_url)
        api_base = f"{parsed_base.scheme}://{parsed_base.netloc}/api/v2/help_center/{locale}"
        sections_url = f"{api_base}/categories/{category_id}/sections.json?per_page=100"
        sections = self._fetch_paginated_items(sections_url, "sections")
        if not sections:
            raise ValueError(f"No sections returned for category {category_id}")

        documents: list[ParsedDocument] = []
        article_count = 0
        for section in sections:
            section_id = section.get("id")
            if not section_id:
                continue
            section_name = normalize_whitespace(section.get("name", ""))
            articles_url = f"{api_base}/sections/{section_id}/articles.json?per_page=100"
            for article in self._fetch_paginated_items(articles_url, "articles"):
                if article.get("draft"):
                    continue
                title = normalize_whitespace(article.get("title") or article.get("name") or "")
                body_text = self._html_to_text(article.get("body") or "")
                if len(body_text) < 80:
                    continue
                documents.append(
                    ParsedDocument(
                        title=title or f"Article {article.get('id')}",
                        content=body_text,
                        source_type="kb_article",
                        source_url=article.get("html_url"),
                        mime_type="text/html",
                        payload_json={
                            "article_id": article.get("id"),
                            "section_id": section_id,
                            "section_name": section_name,
                            "sync_mode": "live_api",
                            "provenance": "zendesk_api",
                            "original_html_url": article.get("html_url"),
                        },
                    )
                )
                article_count += 1
                if article_count >= settings.max_sync_articles:
                    return documents
        if not documents:
            raise ValueError(f"No articles returned for category {category_id}")
        return documents

    def parse_uploaded_document(
        self, filename: str, content_type: str | None, raw_bytes: bytes
    ) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            content = self._read_pdf(raw_bytes)
            mime_type = "application/pdf"
        elif suffix in {".docx", ".doc"}:
            content = self._read_docx(raw_bytes)
            mime_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            content = raw_bytes.decode("utf-8", errors="ignore")
            mime_type = content_type or "text/plain"

        return ParsedDocument(
            title=Path(filename).stem,
            content=normalize_whitespace(content),
            source_type="upload",
            filename=filename,
            mime_type=mime_type,
            payload_json={
                "sync_mode": "upload",
                "provenance": "manager_upload",
            },
        )

    def _read_pdf(self, raw_bytes: bytes) -> str:
        reader = PdfReader(io.BytesIO(raw_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _read_docx(self, raw_bytes: bytes) -> str:
        document = docx.Document(io.BytesIO(raw_bytes))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    def add_document(self, db: Session, revision: AgentRevision, document: ParsedDocument) -> int:
        checksum = sha256_text(f"{document.title}\n{document.content}\n{document.source_url or ''}")
        existing = db.scalar(
            select(KnowledgeDocument).where(
                KnowledgeDocument.revision_id == revision.id,
                KnowledgeDocument.checksum == checksum,
            )
        )
        if existing:
            return 0

        knowledge_document = KnowledgeDocument(
            revision_id=revision.id,
            source_type=document.source_type,
            title=document.title[:500],
            source_url=document.source_url,
            filename=document.filename,
            mime_type=document.mime_type,
            content=document.content,
            checksum=checksum,
            payload_json=dict(document.payload_json or {}),
        )
        db.add(knowledge_document)
        db.flush()

        chunk_count = 0
        metadata_text = normalize_whitespace(
            " ".join(
                [
                    knowledge_document.title,
                    str(knowledge_document.payload_json.get("section_name", "")),
                ]
            )
        )
        for index, chunk in enumerate(chunk_text(document.content)):
            db.add(
                KnowledgeChunk(
                    document_id=knowledge_document.id,
                    revision_id=revision.id,
                    chunk_index=index,
                    content=chunk,
                    token_count=len(tokenize(chunk)),
                    payload_json={
                        "title": knowledge_document.title,
                        "section_name": knowledge_document.payload_json.get("section_name"),
                        "metadata_text": metadata_text,
                        "source_type": knowledge_document.source_type,
                    },
                )
            )
            chunk_count += 1
        return chunk_count

    def sync_revision_sources(
        self,
        db: Session,
        revision: AgentRevision,
        *,
        knowledge_base_url: str | None = None,
        uploaded_documents: list[ParsedDocument] | None = None,
        preserve_corrections: bool = True,
    ) -> SyncOutcome:
        target_url = knowledge_base_url or revision.knowledge_base_url
        revision.sync_status = "syncing"
        db.flush()

        if preserve_corrections:
            docs_to_delete = db.scalars(
                select(KnowledgeDocument).where(
                    KnowledgeDocument.revision_id == revision.id,
                    KnowledgeDocument.source_type != "correction",
                )
            ).all()
        else:
            docs_to_delete = db.scalars(
                select(KnowledgeDocument).where(KnowledgeDocument.revision_id == revision.id)
            ).all()
        for document in docs_to_delete:
            db.delete(document)
        db.flush()

        parsed_documents: list[ParsedDocument] = []
        sync_mode = "fallback"
        fallback_used = False
        last_sync_warning: str | None = None

        if target_url:
            try:
                parsed_documents.extend(self.parse_public_help_center(target_url))
                sync_mode = "live_api"
            except Exception as exc:
                parsed_documents.extend(self._fallback_documents())
                fallback_used = True
                last_sync_warning = (
                    "Live Zendesk KB sync failed; fallback knowledge was loaded instead. "
                    f"Reason: {exc.__class__.__name__}"
                )
        elif not uploaded_documents:
            parsed_documents.extend(self._fallback_documents())
            fallback_used = True
            last_sync_warning = "No knowledge base URL was configured, so fallback knowledge was used."

        if uploaded_documents:
            parsed_documents.extend(uploaded_documents)
            if sync_mode != "live_api":
                sync_mode = "upload"

        if not parsed_documents:
            parsed_documents.extend(self._fallback_documents())
            sync_mode = "fallback"
            fallback_used = True
            last_sync_warning = "No documents were discovered from the source; fallback knowledge was used."

        synced_documents = 0
        synced_chunks = 0
        for parsed in parsed_documents:
            chunk_count = self.add_document(db, revision, parsed)
            if chunk_count == 0:
                continue
            synced_documents += 1
            synced_chunks += chunk_count

        revision.sync_status = "ready"
        if sync_mode == "live_api":
            revision.source_summary = (
                f"Synced {synced_documents} documents and {synced_chunks} chunks "
                "from the live Zendesk Help Center API."
            )
        elif sync_mode == "upload":
            revision.source_summary = (
                f"Synced {synced_documents} documents and {synced_chunks} chunks from uploads."
            )
        else:
            revision.source_summary = (
                f"Loaded {synced_documents} fallback documents and {synced_chunks} chunks."
            )
        payload = dict(revision.payload_json or {})
        payload.update(
            {
                "sync_mode": sync_mode,
                "fallback_used": fallback_used,
                "documents_synced": synced_documents,
                "chunks_synced": synced_chunks,
                "last_sync_warning": last_sync_warning,
                "last_sync_source_url": target_url,
            }
        )
        revision.payload_json = payload
        db.flush()
        return SyncOutcome(
            documents_synced=synced_documents,
            chunks_synced=synced_chunks,
            sync_mode=sync_mode,
            fallback_used=fallback_used,
            last_sync_warning=last_sync_warning,
        )

    def clone_revision(
        self,
        db: Session,
        source_revision: AgentRevision,
        *,
        additional_guidelines: str | None = None,
        knowledge_base_url: str | None = None,
        system_prompt: str | None = None,
        status: str = "published",
    ) -> AgentRevision:
        agent = source_revision.agent
        next_version = (source_revision.version or 0) + 1
        cloned = AgentRevision(
            agent_id=agent.id,
            version=next_version,
            knowledge_base_url=knowledge_base_url
            if knowledge_base_url is not None
            else source_revision.knowledge_base_url,
            additional_guidelines=additional_guidelines
            if additional_guidelines is not None
            else source_revision.additional_guidelines,
            system_prompt=system_prompt if system_prompt is not None else source_revision.system_prompt,
            parent_revision_id=source_revision.id,
            status=status,
            sync_status="idle",
            source_summary=source_revision.source_summary,
            payload_json=dict(source_revision.payload_json or {}),
        )
        db.add(cloned)
        db.flush()

        for document in source_revision.documents:
            cloned_doc = KnowledgeDocument(
                revision_id=cloned.id,
                source_type=document.source_type,
                title=document.title,
                source_url=document.source_url,
                filename=document.filename,
                mime_type=document.mime_type,
                content=document.content,
                checksum=document.checksum,
                payload_json=dict(document.payload_json or {}),
            )
            db.add(cloned_doc)
            db.flush()
            for chunk in document.chunks:
                db.add(
                    KnowledgeChunk(
                        document_id=cloned_doc.id,
                        revision_id=cloned.id,
                        chunk_index=chunk.chunk_index,
                        content=chunk.content,
                        token_count=chunk.token_count,
                        payload_json=dict(chunk.payload_json or {}),
                    )
                )
        db.flush()
        return cloned

    def seed_default_agent(self, db: Session) -> Agent:
        existing = db.scalar(select(Agent).where(Agent.role == "support"))
        if existing:
            return existing

        agent = Agent(
            name="Atome Card Support",
            description="Customer service bot for Atome Card knowledge-base and case-specific mock lookups.",
            role="support",
        )
        db.add(agent)
        db.flush()

        revision = AgentRevision(
            agent_id=agent.id,
            version=1,
            knowledge_base_url=settings.default_kb_url,
            additional_guidelines=(
                "If a customer asks for their own card application status, ask for an application reference if missing and then call the application-status tool. "
                "If a customer asks about their own failed card transaction, ask for a transaction ID if missing and then call the transaction-status tool. "
                "For general informational questions, answer only from verified knowledge-base content with citations."
            ),
            system_prompt=self.build_system_prompt(agent.name),
            published_at=None,
            source_summary="Seeded with fallback knowledge. Run source sync for the latest KB content.",
            payload_json={
                "sync_mode": "fallback",
                "fallback_used": True,
                "documents_synced": len(FALLBACK_KNOWLEDGE),
                "chunks_synced": len(FALLBACK_KNOWLEDGE),
                "last_sync_warning": "Seeded with fallback knowledge until a live KB sync is run.",
                "last_sync_source_url": settings.default_kb_url,
            },
        )
        db.add(revision)
        db.flush()
        agent.active_revision_id = revision.id
        db.flush()

        for item in self._fallback_documents():
            self.add_document(db, revision, item)
        db.flush()
        return agent


source_service = SourceService()
